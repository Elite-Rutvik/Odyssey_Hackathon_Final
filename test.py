import streamlit as st 
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from io import BytesIO
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from docx import Document 

load_dotenv()

genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(BytesIO(pdf.read()))  
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")

def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in the 
    provided context just say, "answer is not available in the context", don't provide the wrong answer.
    Generate questions based on the context only which can contain sub questions in the same topic.
    Don't generate questions that cannot be answered from the given context.
    Context:\n {context}?\n
    Question: \n{question}\n 

    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain

def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)

    docs = new_db.similarity_search(user_question)
    
    chain = get_conversational_chain()

    response = chain(
        {"input_documents": docs, "question": user_question},  
        return_only_outputs=True
    )
    
    answer = response["output_text"]
    st.write("Reply: ", answer)
    
    return user_question, answer

def save_to_word(questions_answers):
    doc = Document()
    doc.add_heading('Question Paper', level=1)
    
    for idx, (question, answer) in enumerate(questions_answers, start=1):
        doc.add_heading(f"Question {idx}:", level=2)
        doc.add_paragraph(question)
        doc.add_heading("Answer:", level=2)
        doc.add_paragraph(answer)
    
    file_path = "Question_Paper.docx"
    doc.save(file_path)
    st.success(f"Document saved as {file_path}")
    st.download_button(
        label="Download Question Paper",
        data=open(file_path, "rb"),
        file_name="Question_Paper.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def main():
    st.set_page_config(page_title="Chat with multiple PDF")
    st.header("Chat with multiple PDF using Gemini")

    questions_answers = []
    user_question = st.text_input("Ask a Question from the PDF files")

    if user_question:
        question, answer = user_input(user_question)
        questions_answers.append((question, answer))
    
    with st.sidebar:
        st.title("Menu:")
        pdf_docs = st.file_uploader("Upload Your PDF files and Click on Submit Button", accept_multiple_files=True, type=["pdf"])
        if st.button("Submit & Proceed"):
            if pdf_docs:
                with st.spinner("Preprocessing..."):
                    raw_text = get_pdf_text(pdf_docs)
                    text_chunks = get_text_chunks(raw_text)
                    get_vector_store(text_chunks)
                    st.success("Done")
            else:
                st.error("Please upload at least one PDF file.")
        
        if st.button("Save to Word"):
            if questions_answers:
                save_to_word(questions_answers)
            else:
                st.error("No questions and answers to save!")

if __name__ == "__main__":
    main()
